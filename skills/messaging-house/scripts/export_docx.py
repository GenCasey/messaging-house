#!/usr/bin/env python3
"""Export a messaging-house markdown file to a black-on-white Word document.

Usage: python export_docx.py messaging-house.md [output.docx]

Handles the markdown this skill produces: #/##/### headings, paragraphs,
bulleted and numbered lists, pipe tables, **bold** and *italic* runs.
Requires python-docx (pip install python-docx).
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is required: pip install python-docx")

BLACK = RGBColor(0, 0, 0)
INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_runs(paragraph, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.color.rgb = BLACK


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.color.rgb = BLACK
        style.font.name = "Calibri"


def flush_table(doc, rows):
    rows = [r for r in rows if not re.match(r"^\|?\s*:?-{2,}", r)]
    if not rows:
        return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    width = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(cells):
        for j in range(width):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs(para, text)
            if i == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path):
    doc = Document()
    style_document(doc)
    table_buffer = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_heading(level=0)
            add_runs(p, line[2:])
        elif line.startswith("## "):
            p = doc.add_heading(level=1)
            add_runs(p, line[3:])
        elif line.startswith("### "):
            p = doc.add_heading(level=2)
            add_runs(p, line[4:])
        elif re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
    if table_buffer:
        flush_table(doc, table_buffer)
    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")
    convert(md_path, out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()

